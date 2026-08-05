from docx import Document
from reportlab.pdfgen import canvas

from app.services.resume_parser import UnsupportedResumeFormat, extract_resume_text


def test_extract_txt(tmp_path):
    path = tmp_path / "resume.txt"
    path.write_text("Jane Doe\nSoftware Engineer at Acme Corp")

    text = extract_resume_text(path)

    assert "Jane Doe" in text
    assert "Acme Corp" in text


def test_extract_docx(tmp_path):
    path = tmp_path / "resume.docx"
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Software Engineer at Acme Corp")
    doc.save(str(path))

    text = extract_resume_text(path)

    assert "Jane Doe" in text
    assert "Acme Corp" in text


def test_extract_pdf(tmp_path):
    path = tmp_path / "resume.pdf"
    c = canvas.Canvas(str(path))
    c.drawString(100, 750, "Jane Doe")
    c.drawString(100, 730, "Software Engineer at Acme Corp")
    c.save()

    text = extract_resume_text(path)

    assert "Jane Doe" in text
    assert "Acme Corp" in text


def test_unsupported_format_raises(tmp_path):
    path = tmp_path / "resume.xyz"
    path.write_text("nonsense")

    try:
        extract_resume_text(path)
        assert False, "expected UnsupportedResumeFormat"
    except UnsupportedResumeFormat:
        pass
